"""
DOCX Document Parser implementation.
Extracts paragraphs and headings from Word documents.
"""

import hashlib
import io
from packages.rag.ingestion.base import BaseDocumentParser, ParsedDocument

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class DOCXDocumentParser(BaseDocumentParser):
    """Parser for Microsoft Word (.docx) documents."""

    async def parse(self, file_content: bytes, filename: str) -> ParsedDocument:
        checksum = hashlib.sha256(file_content).hexdigest()
        extracted_text = ""
        metadata = {}

        if HAS_DOCX:
            try:
                doc = docx.Document(io.BytesIO(file_content))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                extracted_text = "\n\n".join(paragraphs)
                metadata["paragraph_count"] = len(paragraphs)
            except Exception as e:
                extracted_text = f"[DOCX Parsing Warning: {str(e)}]\n" + file_content.decode("latin1", errors="ignore")
        else:
            extracted_text = file_content.decode("latin1", errors="ignore")

        return ParsedDocument(
            filename=filename,
            file_type="docx",
            content=extracted_text,
            checksum=checksum,
            metadata=metadata,
        )
